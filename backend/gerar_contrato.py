import json
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from openai import OpenAI
from dotenv import load_dotenv

# Carrega variáveis de ambiente do arquivo .env
load_dotenv()

# =====================================================================
# 1) EXTRAÇÃO DO PRE-CONTRATO (MESMO DO GEMINI)
# =====================================================================
def extract_contract_data(path):
    doc = Document(path)
    data = {"text": "", "tables": []}

    for p in doc.paragraphs:
        data["text"] += p.text + "\n"

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        data["tables"].append(rows)

    return data


# =====================================================================
# 2) LIMPA MARCAÇÕES
# =====================================================================
def limpa_marcacoes(texto):
    return texto.replace("**", "").replace("--", "—")


# =====================================================================
# 3) SEPARA ASSINATURAS (IGUAL GEMINI)
# =====================================================================
def separar_assinaturas(texto):
    padrao = re.compile(
        r'<<<ASSINATURAS_INICIO>>>(.*?)<<<ASSINATURAS_FIM>>>',
        flags=re.DOTALL | re.IGNORECASE
    )
    m = padrao.search(texto)

    if not m:
        return texto.strip(), ""

    assinaturas = m.group(1).strip()
    corpo = (texto[:m.start()] + texto[m.end():]).strip()
    return corpo, assinaturas


# =====================================================================
# 4) INSERE PARÁGRAFOS
# =====================================================================
def add_paragrafos(doc, texto):
    padrao_clausula = re.compile(r'^CLÁUSULA\s+[A-ZÀ-Ú]+\s*[–—-]\s*.+', re.IGNORECASE)
    padrao_paragrafo = re.compile(r'^PARÁGRAFO\s+[A-ZÀ-Ú]+[:.]?', re.IGNORECASE)

    for line in texto.split("\n"):
        line = line.rstrip()

        if not line:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph("")

        # Cláusula com título em negrito
        if padrao_clausula.match(line):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)

        # Parágrafo com título em negrito
        elif padrao_paragrafo.match(line):
            match = padrao_paragrafo.match(line)
            titulo = match.group(0)
            resto = line[len(titulo):].strip()

            rt = p.add_run(titulo + " ")
            rt.bold = True
            rt.font.size = Pt(12)

            if resto:
                rr = p.add_run(resto)
                rr.font.size = Pt(12)

        # Texto normal
        else:
            run = p.add_run(line)
            run.font.size = Pt(12)

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# =====================================================================
# 5) GERA CONTEÚDO — AGORA ADAPTADO AO BACKEND ATUAL
# =====================================================================
def gerar_conteudo(
    pre_contrato_path: str | None,
    tipo_contrato: str,
    saida_path: str,
    paragrafos_extra=None,
    extra_text: str = "",
    text_area_precontrato: str = ""  
):
    if paragrafos_extra is None:
        paragrafos_extra = []

    BASE_DIR = Path(__file__).resolve().parent

    modelos = {
        "compra-venda": "compra-venda.docx",
        "financiamento-go": "financiamento-go.docx",
        "financiamento-ms": "financiamento-ms.docx",
    }

    if tipo_contrato not in modelos:
        raise ValueError(f"Tipo de contrato inválido: {tipo_contrato}")

    modelo_layout_path = BASE_DIR / modelos[tipo_contrato]
    if not modelo_layout_path.exists():
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_layout_path}")

    # ============================================================
    # 1) O pré-contrato pode vir de arquivo OU do textarea
    # ============================================================

    if text_area_precontrato.strip():
        print("📄 Pré-contrato recebido via TEXTAREA, ignorando .docx")
        dados_extraidos = {
            "text": text_area_precontrato,
            "tables": []  # não há tabelas no texto
        }
    else:
        print("📄 Extraindo dados do pré-contrato .docx...")
        dados_extraidos = extract_contract_data(pre_contrato_path)

    dados_json = json.dumps(dados_extraidos, ensure_ascii=False)

    # ============================================================
    # 2) Config Gemini
    # ============================================================

    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    model = genai.GenerativeModel(
        model_name='gemini-flash-latest',
        generation_config={
            'temperature': 0,
            'top_p': 0.95,
            'top_k': 40,
            'max_output_tokens': 8192,
        }
    )

    print("📄 Lendo layout-modelo base...")
    layout_text = "\n".join([p.text for p in Document(modelo_layout_path).paragraphs])

    # ============================================================
    # 3) PROMPT do Gemini (mesmo do código 1 + textarea)
    # ============================================================
    prompt = f"""

Você é um assistente jurídico especializado em contratos imobiliários.

Tarefa:
- Reescreva o contrato completo (sem resumir nem omitir dados).
- Mantenha a mesma estrutura e títulos do layout.
- Use marcação Markdown para formatação (ex: **negrito**, ### títulos).
- Não altere o cabeçalho, numeração de cláusulas nem o rodapé.
- Quando identificar listas ou quadros de dados (ex: Partes, Posse, Honorários, Comissões, Taxas, Despesas), represente-os como blocos de texto simples, com um título de seção e cada item em uma nova linha.
- Ao final, coloque as assinaturas (nomes, CPFs, testemunhas, data e local), sem marcadores.

🔵 REGRAS CRÍTICAS PARA PARTES (VENDEDORES, COMPRADORES)

Múltiplas Partes: Se houver mais de um vendedor ou mais de um comprador, deixe UMA linha em branco para separar os dados de cada pessoa.

Exemplo CORRETO para Partes:

VENDEDOR(ES):
JOÃO DA SILVA, nacionalidade, estado civil, profissão, portador do RG nº XXX e CPF nº YYY, residente e domiciliado na Rua ZZZ, nº 000, Cidade/Estado, doravante denominado(s) VENDEDOR(ES).

COMPRADOR(ES):
GUSTAVO ALEXANDRE TORRES DE MOURA, detentor de 100% de participação, telefone (55) 62 99125-088 e e-mail gustavoatm@gmail.com, nacionalidade, estado civil, profissão, portador do RG nº XXX e CPF nº YYY, residente e domiciliado na Rua ZZZ, nº 000, Cidade/Estado, doravante denominado(s) COMPRADOR(ES).


🔴 REGRAS CRÍTICAS PARA PARCELAS (leia com atenção):

1. ESTRUTURA GERAL:
   O bloco de parcelas deve começar com o título "Valor e forma de Pagamento", seguido pelo valor total e depois cada parcela, sem linhas em branco extras no início.

   **IMPORTANTE:** NÃO deixe linha em branco entre "Valor total do negócio" e "1ª parcela"

2. FORMATO DE CADA PARCELA (CRÍTICO):
   Cada parcela deve seguir EXATAMENTE este formato com 3 linhas (ou 4, se houver observação):

   **Xª parcela**
   **Valor:** R$ XX.XXX,XX - **Data do Pagamento:** [data ou condição]
   **Forma de pagamento:** [descrever forma COMPLETA incluindo banco, agência, conta, titular, CPF, etc. TUDO em uma linha separado por traços]

   Exemplo CORRETO da estrutura COMPLETA:

   Valor e forma de Pagamento
   Valor total do negócio: R$ 208.000,00 (Duzentos e oito mil reais)
   1ª parcela
   Valor: R$ 12.000,00 - Data do Pagamento: Ato de assinatura do presente instrumento
   Forma de pagamento: TED/PIX - Banco Itau - Agência 4459 - Conta Corrente 84234-2 - titular Deyla Flavia Bertolazzo - CPF 370.990.108-16

   2ª parcela
   Valor: R$ 29.600,00 - Data do Pagamento: Ato da assinatura
   Forma de pagamento: TED/PIX - Banco Itau - Agência 4459 - Conta Corrente 84234-2 - titular Deyla Flavia Bertolazzo - CPF 370.990.108-16

3. O QUE NÃO FAZER (erros comuns):
   ❌ NÃO deixe linha em branco entre "Valor total do negócio" e "1ª parcela"
   ❌ NÃO quebre os dados bancários em múltiplas linhas
   ❌ NÃO coloque cada informação bancária em linha separada
   ❌ NÃO use quebras de linha dentro da "Forma de pagamento"

4. O QUE FAZER:
   ✅ Primeira linha: Título da parcela (ex: "1ª parcela")
   ✅ Segunda linha: Valor e Data juntos (separados por " - ")
   ✅ Terceira linha: "Forma de pagamento: " seguido de TODOS os dados bancários em sequência (separados por " - ")
   ✅ Deixe UMA linha em branco APENAS entre parcelas diferentes (não antes da primeira)

5. TRATAMENTO DE OBSERVAÇÕES/CONDIÇÕES:
   Se houver observações ou condições adicionais da parcela (ex: "FGTS será utilizado", "Financiamento bancário"),
   adicione como quarta linha "Observação: [texto]"

   Exemplo:
   3ª parcela
   Valor: R$ 166.400,00 - Data do Pagamento: Dentro de 120 dias
   Forma de pagamento: Financiamento bancário junto ao banco XYZ
   Observação: Sujeito a aprovação de crédito

🟡 REGRAS CRÍTICAS PARA HONORÁRIOS/COMISSÕES/TAXAS/DESPESAS:

1. SEMPRE PRESERVE ESTAS INFORMAÇÕES: Se o pré-contrato contiver informações sobre:
   - Honorários advocatícios
   - Comissões de corretagem
   - Taxas administrativas
   - Despesas diversas
   - Custos adicionais
   - Responsabilidades financeiras

   VOCÊ DEVE incluí-las no contrato final, INDEPENDENTE do título usado.

2. FORMATO PARA HONORÁRIOS E SIMILARES:
   Se houver qualquer informação sobre custos adicionais, honorários, comissões ou taxas, represente como um bloco de texto formatado:

   [Título adequado: Honorários | Comissões | Taxas | Despesas | etc.]
   [Primeira informação sobre valor/responsável]
   [Segunda informação]
   [...]

3. DETECÇÃO AUTOMÁTICA:
   - Se encontrar termos como "honorário", "comissão", "taxa", "despesa", "custo", "responsabilidade", "pagamento de"
   - Identifique o contexto e crie um bloco de texto apropriado
   - Use o título mais adequado ao contexto (não invente, use o que está no documento ou um similar)

4. EXEMPLOS DE VARIAÇÕES VÁLIDAS:

   Exemplo 1 - Honorários Advocatícios:
   Honorários Advocatícios
   Valor: R$ 5.000,00 (cinco mil reais)
   Responsável: Compradores
   Pagamento: Até a assinatura da escritura

   Exemplo 2 - Comissão de Corretagem:
   Comissão de Corretagem
   Percentual: 6% sobre o valor total
   Valor: R$ 12.480,00
   Responsável: Vendedor

   Exemplo 3 - Múltiplas Despesas:
   Despesas e Responsabilidades
   ITBI: Por conta do comprador
   Registro: Por conta do comprador
   Honorários advocatícios: R$ 3.000,00 - Vendedor
   Certidões: Por conta do vendedor

5. IMPORTANTE:
   ✅ NUNCA omita informações sobre valores, custos ou responsabilidades financeiras
   ✅ Se não houver título claro, use "Despesas e Responsabilidades" ou similar
   ✅ Preserve TODOS os valores e responsáveis mencionados
   ✅ Se estiver após a seção de parcelas, provavelmente é uma despesa/honorário

LEMBRE-SE:
- O bloco de parcelas deve seguir rigorosamente a formatação dos exemplos.
- Qualquer informação sobre honorários, comissões, taxas ou despesas também deve ser formatada como um bloco de texto simples.
- NUNCA omita informações financeiras do documento original.

TEXTO ADICIONAL DO USUÁRIO (textarea):
{extra_text}

LAYOUT DE REFERÊNCIA:
{layout_text}

INFORMAÇÕES EXTRAÍDAS DO PRÉ-CONTRATO:
{dados_json}
    """

    print("🤖 Gerando texto com Gemini...")
    resposta = model.generate_content(prompt)
    conteudo = limpa_marcacoes(resposta.text.strip())

    corpo, assinaturas = separar_assinaturas(conteudo)

    # limpa cabeçalhos duplicados
    remover = [
        r"INSTRUMENTO\s+PARTICULAR.*",
        r"QUADRO\s+RESUMO"
    ]
    for padrao in remover:
        corpo = re.sub(padrao, "", corpo, flags=re.IGNORECASE)

    corpo = re.sub(r"\n{3,}", "\n\n", corpo).strip()

    print("📝 Inserindo conteúdo no modelo .docx...")

    modelo = Document(modelo_layout_path)

    # ponto de inserção
    insert_index = None
    for i, p in enumerate(modelo.paragraphs):
        if "Quadro Resumo" in (p.text or ""):
            insert_index = i + 1
            break
    if insert_index is None:
        insert_index = len(modelo.paragraphs)

    # remove conteúdo após quadro resumo
    while len(modelo.paragraphs) > insert_index:
        p = modelo.paragraphs[-1]
        p._element.getparent().remove(p._element)

    add_paragrafos(modelo, corpo)

    if assinaturas:
        modelo.add_paragraph("")
        add_paragrafos(modelo, assinaturas)

    # cláusulas adicionais selecionadas
    if paragrafos_extra:
        modelo.add_page_break()
        titulo = modelo.add_paragraph("CLÁUSULAS ADICIONAIS")
        titulo.runs[0].bold = True
        for p in paragrafos_extra:
            modelo.add_paragraph(p)
            modelo.add_paragraph("")

    modelo.save(saida_path)

    print(f"✅ Contrato gerado em: {saida_path}")
