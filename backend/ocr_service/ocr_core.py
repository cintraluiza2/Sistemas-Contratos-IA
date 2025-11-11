import os
import json
import re
import pdfplumber
import pytesseract
import docx
from flask import jsonify
from crewai import Crew, Agent, Task, Process
from langchain_openai import ChatOpenAI
from uuid import uuid4
from flask import request

os.environ["OPENAI_API_KEY"] = "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"

from pathlib import Path

OUTPUT_DIR = Path.cwd() / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)



# ---------- Utils ----------
def extract_page_text(page, langs: str = "por+eng") -> str:
    """Extrai texto da página; se tiver pouco texto, faz OCR."""
    text = (page.extract_text() or "").strip()
    if len(text) >= 30:
        return text
    pil_image = page.to_image(resolution=216).original
    ocr_text = pytesseract.image_to_string(pil_image, lang=langs)
    return (ocr_text or "").strip()

def normalize_text(s: str) -> str:
    s = s.replace("\r\n", "\n")
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def extract_docx_text(file_storage) -> str:
    """Extrai texto de .docx usando python-docx."""
    try:
        file_storage.stream.seek(0)
        d = docx.Document(file_storage.stream)
        parts = []
        for p in d.paragraphs:
            parts.append(p.text or "")
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return normalize_text("\n".join(parts))
    except Exception as e:
        print(f"[DOCX] erro: {e}")
        return ""

def build_documents_from_request(req_files, langs="por+eng"):
    """Processa múltiplos arquivos e retorna (nomes, textos)."""
    textos, nomes = [], []
    for f in req_files:
        nome = (f.filename or "documento").strip()
        low = nome.lower()
        nomes.append(nome)
        try:
            if low.endswith(".pdf"):
                f.stream.seek(0)
                with pdfplumber.open(f.stream) as pdf:
                    partes = []
                    for i, page in enumerate(pdf.pages, start=1):
                        partes.append(f"[PÁGINA {i}]\n{extract_page_text(page, langs)}")
                    conteudo = normalize_text("\n\n".join(partes))
                textos.append(f"### {nome}\n{conteudo}")
            elif low.endswith(".docx"):
                conteudo = extract_docx_text(f)
                textos.append(f"### {nome}\n{conteudo or '(sem texto extraído do DOCX)'}")
            else:
                f.stream.seek(0)
                raw = f.stream.read()
                if isinstance(raw, bytes):
                    raw = raw.decode("utf-8", errors="ignore")
                conteudo = normalize_text(str(raw or ""))
                textos.append(f"### {nome}\n{conteudo}")
        except Exception as e:
            print(f"[OCR] Erro ao processar {nome}: {e}")
    return nomes, textos

def require_api_key_or_500():
    if not os.getenv("OPENAI_API_KEY"):
        return jsonify({"error": "OPENAI_API_KEY ausente"}), 500
    return None

# ---------- CrewAI Agents & Tasks ----------

# LLMs
llm_analise_distribuicao = ChatOpenAI(
    model="gpt-4o",
    temperature=0.3,
    max_retries=2,
    timeout=120  
)

llm_resumo_executivo = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.5,
    max_retries=2,
    timeout=60
)

# Agentes
agente_analista_tecnico = Agent(
    role = "Advogado Especialista em Análise de Documentos Pessoais (RG, CPF, Certidões de Óbito e de Casamento)",
    goal = """Emitir pareceres jurídicos claros e fundamentados sobre os riscos,
              impactos legais e consequências da ausência, irregularidade ou 
              inconsistência de documentos pessoais em diferentes contextos
              (civil, trabalhista, administrativo e empresarial).""",
    backstory = """Você é um advogado sênior com ampla experiência em direito civil 
                  e registral, consultor em conformidade documental e assessor de 
                  empresas e cidadãos. Seu ponto forte é interpretar a legislação 
                  aplicável, identificar riscos jurídicos e propor soluções práticas 
                  para evitar litígios e garantir segurança jurídica.""",
    verbose=True,
    allow_delegation=False,
    llm=llm_analise_distribuicao,
)

agente_resumidor_executivo = Agent(
    role='Especialista em Comunicação e Síntese Técnica',
    goal="""Transformar análises técnicas detalhadas em resumos concisos e claros,
         utilizando formato de bullet points, para facilitar a compreensão por
         públicos diversos (executivos, outras áreas).""",
    backstory="""Você tem a habilidade única de pegar informações técnicas densas
               e extrair a essência, comunicando os pontos mais importantes de
               forma direta e acionável. Seu foco é a clareza e a concisão.""",
    verbose=True,
    allow_delegation=False,
    llm=llm_resumo_executivo,
)

# Tarefas
tarefa_analise_tecnica = Task(
description = f"""
Analise detalhadamente o seguinte conjunto de documentos pessoais fornecido:
--- INÍCIO DOS DOCUMENTOS ---
    {{regulamento_texto}}
--- FIM DOS DOCUMENTOS ---

Sua tarefa é analisar documentos fornecidos pela imobiliária *My Broker* 
(pré-contrato de compra e venda, matrícula do imóvel, certidões de casamento/divórcio/óbito, 
comprovantes de endereço, certidões fiscais, etc.) e produzir um *parecer jurídico completo, 
didático e consultivo*.  

Esse parecer será enviado pela *Batista Advogados Associados* à *My Broker*, de forma 
profissional e clara, para orientar sobre os riscos e soluções antes da formalização do contrato.

## Estrutura obrigatória do parecer:

1. IDENTIFICAÇÃO
   - Assunto: [descrever]
   - Interessada: Imobiliária My Broker
   - Proprietário(s) analisado(s): [nome(s) do(s) vendedor(es)]
   - Imóvel: [matrícula ]
   - Data: [data da análise]
   

2. CONTEXTO
   - Resumo dos documentos analisados.
   - Objetivo da análise: identificar inconsistências, riscos e propor soluções práticas.

3. ANÁLISE DE INCONSISTÊNCIAS E RISCOS
   - Liste cada inconsistência em tópicos numerados.
   - Para cada inconsistência, detalhe:
     - *Descrição* do problema.
     - *Gravidade* (baixa, média ou alta).
   - Exemplos: divórcio não averbado, partilha incompleta, débitos de IPTU, divergência em dados da matrícula, ausência de documentos, duplicidade e/ou ausência de CPF/CNPJ.

4. IMPLICAÇÕES LEGAIS
   - Explique, em linguagem clara, o que pode acontecer se cada inconsistência não for corrigida:
     - risco de contestação judicial,
     - negativa de registro da escritura,
     - penhora do imóvel,
     - desgaste para a My Broker na intermediação.

5. SOLUÇÕES PROPOSTAS
   - Para cada inconsistência, indique:
     - *O que deve ser feito* (averbar divórcio, apresentar formal de partilha, quitar IPTU, etc.).
     - *Quem deve providenciar* (vendedor, herdeiros, etc.).
     - *Quando* (antes da escritura, antes do registro, etc.).
     - *Documentos necessários*.
     - *Prazo médio* quando aplicável.

6. INSTRUÇÕES E RECOMENDAÇÕES
   - Liste em formato passo a passo (Passo 1, Passo 2, Passo 3).  
   - Indique local de obtenção dos documentos (cartório de registro civil, cartório de imóveis, prefeitura, etc.).  
   - Informe *impactos diretos para a My Broker* caso não haja regularização: atrasos, insegurança para compradores, risco de nulidade, perda de credibilidade.

7. CONCLUSÃO
   - Reforce que as pendências podem ser regularizadas.  
   - Explique de forma realista se é possível ou não prosseguir com a venda antes da regularização, e quais riscos a My Broker assume.  
   - Indique a posição final do *Batista Advogados Associados* sobre a operação.
Sua análise deve ser minuciosa, clara e enriquecida com referências a normas jurídicas brasileiras.
""",
expected_output = """
Um parecer jurídico detalhado e bem fundamentado,
cobrindo todos os aspectos solicitados na descrição.
O parecer deve indicar claramente quando informações
adicionais da busca web foram utilizadas para enriquecer a análise
(ex.: 'Segundo o art. 215 do Código Civil...', 'De acordo com a Lei de Registros Públicos...').
""",
    agent=agente_analista_tecnico
)

tarefa_resumo_bulletpoints = Task(
    description=f"""
    Você recebeu um relatório técnico detalhado sobre um novo regulamento.
    Sua tarefa é ler este relatório e sintetizar as informações mais cruciais
    em um formato de bullet points (lista de marcadores).

    Crie uma lista concisa (idealmente entre 5 a 8 bullet points) que capture:
    - A identificação, contexto, análise de inconsistências e riscos, implifcações legais, soluções propostas, instruções e
    recomendações e conclusão.

    O resumo deve ser claro, direto e fácil de entender por não especialistas.
    """,
    expected_output="""Uma lista de 5 a 8 bullet points concisos e claros,
                     resumindo os aspectos mais importantes da análise técnica
                     do regulamento.""",
    agent=agente_resumidor_executivo,
    context=[tarefa_analise_tecnica]
)

crew_analise_resumo = Crew(
    agents=[agente_analista_tecnico, agente_resumidor_executivo],
    tasks=[tarefa_analise_tecnica, tarefa_resumo_bulletpoints],
    process=Process.sequential,
    verbose=True,
)

def clean_markdown(text: str) -> str:
    text = re.sub(r"(\*{1,2}|#{1,6})", "", text)
    return text.strip()

# ---------- Core handler ----------
def handle_analisar():
    maybe_err = require_api_key_or_500()
    if maybe_err:
        return maybe_err

    if not request.files:
        return jsonify({"error": "Arquivo não enviado."}), 400

    #  aceita "file" e "files"
    files = request.files.getlist("file")
    if not files:
        files = request.files.getlist("files")

    if not files:
        return jsonify({"error": "Nenhum arquivo válido recebido."}), 400

    nomes, textos = build_documents_from_request(files, langs="por+eng")
    if not textos:
        return jsonify({"error": "Falha ao extrair texto dos arquivos."}), 422

    entrada = "\n\n".join(textos)
    inputs_da_crew = {"regulamento_texto": entrada}
    crew_analise_resumo.kickoff(inputs=inputs_da_crew)

    parecer_detalhado = tarefa_analise_tecnica.output.raw
    resumo_exec = tarefa_resumo_bulletpoints.output.raw

    #  Função para limpar Markdown e símbolos
    def clean_output(text: str) -> str:
        if not text:
            return ""
        # 1) bold/italics markdown: **texto** ou __texto__
        text = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text, flags=re.DOTALL)
        # 2) cabeçalhos ##, ### etc. no início da linha
        text = re.sub(r'(?m)^\s{0,3}#{1,6}\s*', '', text)
        # 3) linhas de separador --- (ou mais)
        text = re.sub(r'(?m)^\s*-{3,}\s*$', '', text)
        # 4) traços duplos ( -- ou --- no meio da frase) -> espaço
        text = re.sub(r'\s-{2,}\s', ' ', text)
        # 5) bullets especiais -> hífen simples (mantemos um hífen só)
        text = re.sub(r'[•●▪▶►]', '-', text)
        # 6) espaços/linhas
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 7) aspas tipográficas
        text = (text.replace('“','"').replace('”','"').replace('’',"'"))
        text = re.sub(r'\*', '', text)  # remove asteriscos soltos
        return text.strip()

    #  Limpa os textos do LLM
    parecer_detalhado = clean_output(parecer_detalhado)
    resumo_exec = clean_output(resumo_exec)

    case_id = uuid4().hex
    parecer_path = OUTPUT_DIR / f"{case_id}_parecer.json"
    with open(parecer_path, "w", encoding="utf-8") as f:
        json.dump({
            "analise_tecnica": parecer_detalhado,
            "resumo": resumo_exec,
            "documentos": nomes,
        }, f, ensure_ascii=False, indent=2)

    return jsonify({
        "ok": True,
        "case_id": case_id,
        "documentos": nomes,
        "analise_tecnica": parecer_detalhado,
        "resumo": resumo_exec,
        "missingData": [],
        "inconsistencies": [],
        "cpfDuplicates": [],
        "risks": [],
        "safePoints": [],
        "usedOCR": False
    }), 200
